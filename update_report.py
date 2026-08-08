import docx
from docx.shared import Pt, RGBColor
from docx.oxml import OxmlElement

doc = docx.Document('ShopEase_Finalt.docx')

# 1. Inject Expanded Literature Review
# Find "2.2 Literature Review"
lit_rev_idx = -1
for i, p in enumerate(doc.paragraphs):
    if p.text.strip() == '2.2 Literature Review':
        lit_rev_idx = i
        break

if lit_rev_idx != -1 and (lit_rev_idx + 1) < len(doc.paragraphs):
    p = doc.paragraphs[lit_rev_idx + 1]
    p.text = (
        'Recent studies in modern web application development highlight a significant shift towards lightweight, component-driven architectures. '
        'Traditional e-commerce platforms often relied on heavy MVC (Model-View-Controller) structures and monolithic databases, which hindered rapid feature deployment and mobile responsiveness. '
        'According to recent software engineering trends, utilizing utility-first CSS frameworks (such as Tailwind CSS) drastically improves rendering speeds by minimizing unused CSS. '
        'Furthermore, research into client-side state management demonstrates that leveraging native solutions like the React Context API is often more efficient and less boilerplate-heavy than introducing third-party libraries like Redux for mid-scale applications.\n\n'
        'In the domain of real-time data handling, the adoption of NoSQL database solutions (such as Firebase Firestore) has proven to be highly effective. '
        'Unlike traditional SQL databases, NoSQL document-based models allow for instantaneous data synchronization across connected clients, which is crucial for modern e-commerce features like live cart updates and inventory tracking. '
        'ShopEase Nepal integrates these modern methodologies by combining React, Tailwind CSS, and Firebase to deliver a robust, highly responsive, and real-time e-commerce experience.'
    )

# 2. Inject System Algorithms
# Find "Chapter 4" to insert right before it
chap4_idx = -1
for i, p in enumerate(doc.paragraphs):
    if 'Chapter 4' in p.text and 'IMPLEMENTATION' in p.text:
        chap4_idx = i
        break

if chap4_idx != -1:
    # Insert new paragraphs before Chapter 4
    # We do this by inserting before the Chapter 4 paragraph element
    chap4_p = doc.paragraphs[chap4_idx]
    
    # We use insert_paragraph_before
    h = chap4_p.insert_paragraph_before('3.6 System Algorithms', style='Heading 2')
    
    p1 = chap4_p.insert_paragraph_before('ShopEase Nepal employs specific algorithms to optimize the user experience in product discovery and catalog management:', style='Normal')
    
    # Algorithm 1
    h_alg1 = chap4_p.insert_paragraph_before('1) Content-Based Recommendation Algorithm', style='Heading 3')
    chap4_p.insert_paragraph_before('Purpose:\nThis algorithm suggests related products to users based on the category and tags of the item they are currently viewing.', style='Normal')
    chap4_p.insert_paragraph_before('Process:\na) Extract the category and tags of the current product.\nb) Query the database for other products that share the same category.\nc) Calculate a similarity score based on matching tags.\nd) Rank the products and display the top results in a "You might also like" section.', style='Normal')
    chap4_p.insert_paragraph_before('Result:\nIncreases user engagement and cross-selling by dynamically suggesting relevant items without requiring complex machine learning models.', style='Normal')
    
    # Algorithm 2
    h_alg2 = chap4_p.insert_paragraph_before('2) Linear Search & Filter Algorithm', style='Heading 3')
    chap4_p.insert_paragraph_before('Purpose:\nQuickly filters the product catalog on the frontend to provide instant feedback to the user without overwhelming the backend database with continuous queries.', style='Normal')
    chap4_p.insert_paragraph_before('Process:\na) The entire product catalog is cached in the local React Context state upon initial load.\nb) When a user enters a search query, the algorithm iterates over the product array.\nc) It applies a case-insensitive regular expression match against the product name and description.\nd) Excludes any products that do not belong to the currently selected category filter.', style='Normal')
    chap4_p.insert_paragraph_before('Result:\nDelivers a highly responsive, zero-latency search experience for the end-user while minimizing database read operations and costs.', style='Normal')
    

# 3. Clean empty gaps and 4. Force Font Color to Black
paragraphs_to_remove = []

for p in doc.paragraphs:
    text = p.text.strip()
    
    # Check if empty
    if not text:
        # Check if it has a page break or something important
        has_break = False
        for r in p.runs:
            if 'w:br' in r._element.xml or 'lastRenderedPageBreak' in r._element.xml:
                has_break = True
                break
        if not has_break:
            paragraphs_to_remove.append(p)
            continue # Don't format removed paragraphs

    # Force font color black
    for r in p.runs:
        r.font.color.rgb = RGBColor(0, 0, 0)
        
        # Also ensure style is black if inheriting? No, explicit run color overrides it.

# Remove the empty paragraphs
for p in paragraphs_to_remove:
    p._element.getparent().remove(p._element)

# Also force table colors if any exist
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.color.rgb = RGBColor(0, 0, 0)

doc.save('ShopEase_Finalt.docx')
print('Document updated successfully.')
