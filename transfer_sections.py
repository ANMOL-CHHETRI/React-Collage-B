import docx
from docx.shared import Pt, RGBColor, Inches
import re
import os

doc_in = docx.Document('ShopEase_Documentation.docx')
doc_out = docx.Document('ShopEase_Finalt.docx')

# Helper to find where to insert in doc_out
def find_insert_point(doc, marker_text):
    for i, p in enumerate(doc.paragraphs):
        if marker_text in p.text:
            return i
    return -1

def extract_section(start_title, end_titles):
    section_data = [] # stores dicts {'type': 'text'|'image', 'data': ..., 'style': ...}
    recording = False
    
    for p in doc_in.paragraphs:
        text = p.text.strip()
        
        # Check if we hit the start title
        if start_title in text:
            recording = True
            
        # Check if we hit an end title
        for end_t in end_titles:
            if end_t in text and recording:
                recording = False
                return section_data
                
        if recording:
            # Check for images
            has_image = False
            for r in p.runs:
                if 'pic:pic' in r._element.xml or 'w:drawing' in r._element.xml:
                    # Find relationship ID
                    try:
                        # Very basic regex to find rId
                        rId_match = re.search(r'r:embed="(rId\d+)"', r._element.xml)
                        if rId_match:
                            rId = rId_match.group(1)
                            image_part = doc_in.part.related_parts[rId]
                            image_bytes = image_part.blob
                            # Save to temp file
                            temp_name = f'temp_{rId}.png'
                            with open(temp_name, 'wb') as f:
                                f.write(image_bytes)
                            section_data.append({'type': 'image', 'data': temp_name, 'style': p.style.name})
                            has_image = True
                    except Exception as e:
                        print("Error extracting image:", e)
            
            if not has_image and text:
                section_data.append({'type': 'text', 'data': text, 'style': p.style.name})
                
    return section_data

def inject_section(idx, title, section_data, header_level='Heading 2'):
    if idx == -1: return idx
    target_p = doc_out.paragraphs[idx]
    
    # Insert Title
    target_p.insert_paragraph_before(title, style=header_level)
    
    for item in section_data:
        if item['type'] == 'text':
            # don't inject the original title
            if not (item['data'] and item['data'][0].isdigit() and item['data'][1:3] == '. '):
                p = target_p.insert_paragraph_before(item['data'], style=item['style'])
                for r in p.runs:
                    r.font.color.rgb = RGBColor(0, 0, 0)
        elif item['type'] == 'image':
            p = target_p.insert_paragraph_before('')
            run = p.add_run()
            try:
                run.add_picture(item['data'], width=Inches(6.0))
            except Exception as e:
                print("Failed to add picture:", e)
    return idx

# 1. Chapter 3 Additions (14, 15, 16) -> Insert before Chapter 4
idx_chap4 = find_insert_point(doc_out, 'Chapter 4')

uc_data = extract_section('14. Use Case Diagram', ['15. Data Flow', '16. Entity', '17. Conclusion'])
inject_section(idx_chap4, '3.7 Use Case Diagram', uc_data)

dfd_data = extract_section('15. Data Flow', ['16. Entity', '17. Conclusion'])
inject_section(idx_chap4, '3.8 Data Flow Diagrams (DFD)', dfd_data)

erd_data = extract_section('16. Entity Relationship', ['17. Conclusion'])
inject_section(idx_chap4, '3.9 Entity Relationship Diagram (ERD)', erd_data)

# 2. Chapter 4 Additions (10, 11) -> Insert before Chapter 5
idx_chap5 = find_insert_point(doc_out, 'Chapter 5')

auth_data = extract_section('10. Authentication System', ['11. Deployment', '12. Future', '13. System UI'])
inject_section(idx_chap5, '4.4 Authentication System', auth_data)

deploy_data = extract_section('11. Deployment', ['12. Future', '13. System UI'])
inject_section(idx_chap5, '4.5 Deployment', deploy_data)

# 3. Chapter 5 Additions (12) -> Insert before Appendix/List of Figures
idx_end = find_insert_point(doc_out, 'List of Figures')
if idx_end == -1: idx_end = len(doc_out.paragraphs) - 1

future_data = extract_section('12. Future Enhancements', ['13. System UI', '14. Use Case', '17. Conclusion'])
# Find 5.3 to insert there? Or just append to end of Chap 5
# Actually inserting before 'List of Figures' puts it at end of Chap 5.
inject_section(idx_end, '5.4 Additional Enhancements', future_data)

# 4. Appendix Additions (13. System UI Screenshots)
ui_data = extract_section('13. System UI Screenshots', ['14. Use Case', '15. Data Flow', '16. Entity', '17. Conclusion'])
# Append at the very end
doc_out.add_heading('Appendix A: System UI Screenshots', level=1)
for item in ui_data:
    if item['type'] == 'text':
        if not (item['data'] and item['data'][0].isdigit() and item['data'][1:3] == '. '):
            p = doc_out.add_paragraph(item['data'], style=item['style'])
            for r in p.runs:
                r.font.color.rgb = RGBColor(0, 0, 0)
    elif item['type'] == 'image':
        p = doc_out.add_paragraph('')
        run = p.add_run()
        try:
            run.add_picture(item['data'], width=Inches(6.0))
        except:
            pass

# Ensure fonts are black everywhere just in case
for p in doc_out.paragraphs:
    for r in p.runs:
        r.font.color.rgb = RGBColor(0, 0, 0)

doc_out.save('ShopEase_Finalt.docx')
print("Successfully merged missing sections and images!")

# Cleanup temp images
for file in os.listdir('.'):
    if file.startswith('temp_') and file.endswith('.png'):
        try:
            os.remove(file)
        except:
            pass
