from docx import Document
from docx.shared import Pt, RGBColor

weeks_data = [
    {
        'week': 1,
        'phase': 'Project Initialization & Core Routing',
        'progress': '15% Complete',
        'tasks': [
            'Project Setup: Initialized React project with Vite, setting up the foundational environment.',
            'Routing Basics: Configured initial routes for HomePage and AboutPage.',
            'Navigation: Created basic Header and Navbar components.',
            'Cleanup: Removed unnecessary boilerplate code (index.css, old App.jsx contents).'
        ],
        'challenges': [
            'Configuration: Setting up initial project structure and dependencies.',
            'Boilerplate: Removing unwanted parts without breaking the Vite setup.'
        ],
        'solutions': [
            'Clean Slate: Established a clean starting point by methodically clearing out default Vite assets.',
            'Component Structure: Created a structured src/components directory from day one.'
        ],
        'contribs': [
            ('Anmol Chhetri (Lead)', 'Project Oversight', 'Project setup and initialization.'),
            ('Sahil Tuladhar', 'Backend Developer', 'Configured initial routing.'),
            ('Sarang Limbu', 'Frontend Developer', 'Created initial Header component.'),
            ('Sanskriti Maharjan', 'UI Component Developer', 'Basic CSS setup and cleanup.'),
            ('Smriti Tamang', 'QA & Testing', 'Initial project build verification.')
        ]
    },
    {
        'week': 2,
        'phase': 'Dashboard Layouts & Navigation',
        'progress': '30% Complete',
        'tasks': [
            'Dashboards: Developed basic UI scaffolding for Admin and User Dashboards.',
            'Navigation: Updated links and structure in Header/Navbar for better flow.',
            'Styling: Implemented responsive CSS updates across the application.',
            'Consistency: Standardized basic UI elements.'
        ],
        'challenges': [
            'Layout Consistency: Ensuring consistent styling across different dashboard views.',
            'Responsiveness: Making the initial Navbar work well on different screen sizes.'
        ],
        'solutions': [
            'Shared Classes: Centralized CSS classes and established a common layout structure.',
            'Flexbox/Grid: Utilized modern CSS layout techniques for responsive design.'
        ],
        'contribs': [
            ('Anmol Chhetri (Lead)', 'Project Oversight', 'Coordinated dashboard requirements.'),
            ('Sahil Tuladhar', 'Backend Developer', 'Admin Dashboard scaffolding.'),
            ('Sarang Limbu', 'Frontend Developer', 'Navbar enhancements and responsiveness.'),
            ('Sanskriti Maharjan', 'UI Component Developer', 'User Dashboard initial UI design.'),
            ('Smriti Tamang', 'QA & Testing', 'Navigation link testing.')
        ]
    },
    {
        'week': 3,
        'phase': 'Authentication UI & Documentation',
        'progress': '45% Complete',
        'tasks': [
            'Login UI: Enhanced User Login page UI with a visually appealing banner image.',
            'Documentation: Drafted initial project documentation and README files.',
            'Team Structure: Finalized development team roles and documented them.',
            'Forms: Set up basic input structures for authentication.'
        ],
        'challenges': [
            'Design Aesthetics: Designing an attractive login page that fits the premium e-commerce theme.',
            'Documentation: Ensuring all team roles and initial project specs were accurately captured.'
        ],
        'solutions': [
            'Visual Assets: Added a high-quality banner image and refined the form layout for better user experience.',
            'Markdown: Created structured markdown files to track progress and team assignments.'
        ],
        'contribs': [
            ('Anmol Chhetri (Lead)', 'Project Oversight', 'Drafted initial project documentation.'),
            ('Sahil Tuladhar', 'Backend Developer', 'Prepared forms for future login state management.'),
            ('Sarang Limbu', 'Frontend Developer', 'Implemented Login UI layout.'),
            ('Sanskriti Maharjan', 'UI Component Developer', 'Integrated banner and form styling.'),
            ('Smriti Tamang', 'QA & Testing', 'Form layout and responsiveness testing.')
        ]
    },
    {
        'week': 4,
        'phase': 'Theming, Maps, & QA Framework',
        'progress': '60% Complete',
        'tasks': [
            'Dark Mode: Implemented dark mode variants across the application.',
            'Map Integration: Added and fixed bugs related to the delivery coverage map.',
            'QA Setup: Added formal QA bug report templates and processes.',
            'FAQ Page: Implemented the FAQ Page UI.'
        ],
        'challenges': [
            'State Management: Ensuring dark mode applies consistently across all components.',
            'Version Control: Resolving merge conflicts resulting from parallel development tracks.'
        ],
        'solutions': [
            'Context API: Used React Context (AuthContext) to manage global theme state.',
            'Communication: Improved team coordination to reduce complex merge conflicts.'
        ],
        'contribs': [
            ('Anmol Chhetri (Lead)', 'Project Oversight', 'Merge conflict resolution and team coordination.'),
            ('Sahil Tuladhar', 'Backend Developer', 'Map functionality integration and fixes.'),
            ('Sarang Limbu', 'Frontend Developer', 'Developed the FAQ Page UI.'),
            ('Sanskriti Maharjan', 'UI Component Developer', 'Implemented global dark mode styles.'),
            ('Smriti Tamang', 'QA & Testing', 'Created QA bug report templates.')
        ]
    },
    {
        'week': 5,
        'phase': 'Product Catalog & User Profile',
        'progress': '75% Complete',
        'tasks': [
            'Product Data: Added Traditional Apparel products and improved sorting logic.',
            'User Profile: Redesigned the User Profile UI for better usability.',
            'Footer: Created a global footer component and added it to multiple pages.',
            'Image Optimization: Addressed image loading and referrer policy issues.'
        ],
        'challenges': [
            'Security & Performance: Preventing image referrer leakage and optimizing image loading from external CDNs.',
            'Component Reusability: Building a footer that works consistently across all page layouts.'
        ],
        'solutions': [
            'Security Policies: Added referrerPolicy=\'no-referrer\' to img tags.',
            'UX Enhancements: Created an ImageWithSkeleton component for smoother image loading experiences.'
        ],
        'contribs': [
            ('Anmol Chhetri (Lead)', 'Project Oversight', 'Product data management and curation.'),
            ('Sahil Tuladhar', 'Backend Developer', 'Implemented robust product sorting logic.'),
            ('Sarang Limbu', 'Frontend Developer', 'Created and integrated the global footer component.'),
            ('Sanskriti Maharjan', 'UI Component Developer', 'Redesigned the User Profile interface.'),
            ('Smriti Tamang', 'QA & Testing', 'Verified image loading behaviors and skeleton animations.')
        ]
    },
    {
        'week': 6,
        'phase': 'Page Refinements & E2E Testing Prep',
        'progress': '85% Complete',
        'tasks': [
            'Page Updates: Refined Contact, UserLogin, ProductDetails, and AdminLogin pages.',
            'Shopping Experience: Enhanced overall user shopping experience and order action flows.',
            'Testing: Added initial Playwright configuration for End-to-End (E2E) testing.',
            'Cleanup: Removed unnecessary images and files from the repository.'
        ],
        'challenges': [
            'Complex Flows: Handling complex user flows and ensuring all pages meet high quality standards.',
            'Testing Setup: Configuring Playwright correctly for a React/Vite environment.'
        ],
        'solutions': [
            'Iterative Polish: Conducted thorough manual testing and iterative UI refinements.',
            'Test Infrastructure: Established the foundational setup for automated E2E tests.'
        ],
        'contribs': [
            ('Anmol Chhetri (Lead)', 'Project Oversight', 'Enhanced order actions and user shopping flows.'),
            ('Sahil Tuladhar', 'Backend Developer', 'Set up initial Playwright configuration.'),
            ('Sarang Limbu', 'Frontend Developer', 'Updated ProductDetails page functionality.'),
            ('Sanskriti Maharjan', 'UI Component Developer', 'Refined Contact page UI.'),
            ('Smriti Tamang', 'QA & Testing', 'E2E test planning and manual page verification.')
        ]
    },
    {
        'week': 7,
        'phase': 'Final UI/UX Polish & Authentication Flows',
        'progress': '95% Complete',
        'tasks': [
            'Sign-up UI Redesign: Implemented a new slide-up modal drawer for user registration.',
            'Login Animation: Added a cinematic store opening animation overlay on login.',
            'Admin Login Revamp: Refactored Admin Login to a clean card layout with no images.',
            'Hidden Access: Created a secret 4-corner tap sequence to unlock admin login securely.',
            'Auth Flow Management: Patched AuthContext to give UserLogin page control over redirect timing.'
        ],
        'challenges': [
            'Animation Sync: Ensuring complex CSS keyframes timed perfectly with React state and route redirects.',
            'Security via Obscurity: Implementing a hidden admin trigger without breaking normal user flows.'
        ],
        'solutions': [
            'Controlled Timeouts: Used setTimeout linked to exact CSS animation durations for routing.',
            'State Separation: Handled sign-up flow through local states inside UserLoginPage rather than global context.'
        ],
        'contribs': [
            ('Anmol Chhetri (Lead)', 'Project Oversight', 'Managed UI/UX transition and reviewed auth workflow logic.'),
            ('Sahil Tuladhar', 'Backend Developer', 'Handled localStorage data structures and state persistence.'),
            ('Sarang Limbu', 'Frontend Developer', 'Developed core React state logic for sign-up modal and Easter eggs.'),
            ('Sanskriti Maharjan', 'UI Component Developer', 'Designed and implemented custom CSS keyframe animations.'),
            ('Smriti Tamang', 'QA & Testing', 'Verified visual timing of animations and tested the secret admin trigger.')
        ]
    }
]

for w in weeks_data:
    doc = Document()
    doc.add_heading(f"WEEKLY PROGRESS LOG - WEEK {w['week']}", 0)
    
    p = doc.add_paragraph()
    p.add_run('Project Title: ').bold = True
    p.add_run('ShopEase Nepal (Ecommerce Platform)\n')
    p.add_run('Current Phase: ').bold = True
    p.add_run(f"{w['phase']}\n")
    p.add_run('Overall Progress: ').bold = True
    p.add_run(w['progress'])

    doc.add_heading('1. TASKS COMPLETED', level=1)
    for task in w['tasks']:
        doc.add_paragraph(task, style='List Bullet')

    doc.add_heading('2. CHALLENGES', level=1)
    for c in w['challenges']:
        doc.add_paragraph(c, style='List Bullet')

    doc.add_heading('3. SOLUTIONS', level=1)
    for s in w['solutions']:
        doc.add_paragraph(s, style='List Bullet')

    doc.add_heading('4. INDIVIDUAL TASK CONTRIBUTIONS', level=1)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Member'
    hdr_cells[1].text = 'Focus Area'
    hdr_cells[2].text = 'Key Deliverables'

    for item in w['contribs']:
        row_cells = table.add_row().cells
        row_cells[0].text = item[0]
        row_cells[1].text = item[1]
        row_cells[2].text = item[2]

    doc.add_heading('5. LEADERSHIP REVIEW & STATUS REPORT', level=1)
    
    if w['week'] < 7:
        doc.add_paragraph(f"The project successfully completed Week {w['week']}, reaching {w['progress']}. The team focused heavily on {w['phase'].lower()}, resolving key challenges along the way.")
    else:
        doc.add_paragraph('The project has reached the 95% completion mark. With the core data structures and routing already established, the team has successfully shifted focus towards premium UI/UX features and complex CSS animations. The platform now supports a highly polished authentication flow, making it ready for the final integration of the actual backend API.')

    p2 = doc.add_paragraph()
    p2.add_run('Current Status: ').bold = True
    p2.add_run(f"{w['progress']}\n")
    p2.add_run('Next Milestone: ').bold = True
    if w['week'] < 7:
        p2.add_run(f"Proceed to Week {w['week']+1} objectives.")
    else:
        p2.add_run('Full integration of the REST API backend and final checkout functionality refinements.')

    doc.save(f"WEEKLY_PROGRESS_LOG_WEEK_{w['week']}.docx")
    print(f"Successfully created WEEKLY_PROGRESS_LOG_WEEK_{w['week']}.docx")
