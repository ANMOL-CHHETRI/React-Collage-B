import docx
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc_out = docx.Document()

# Styles
for style in doc_out.styles:
    if hasattr(style, 'font'):
        style.font.name = 'Arial'

h1 = doc_out.styles['Heading 1']
h1.font.size = Pt(18)
h1.font.bold = True
h1.paragraph_format.line_spacing = 1.5
h1.paragraph_format.space_after = Pt(8)
h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
h1.paragraph_format.page_break_before = True

try:
    h2 = doc_out.styles['Heading 2']
except KeyError:
    h2 = doc_out.styles.add_style('Heading 2', docx.enum.style.WD_STYLE_TYPE.PARAGRAPH)
h2.font.size = Pt(14)
h2.font.bold = True
h2.paragraph_format.line_spacing = 1.5
h2.paragraph_format.space_after = Pt(8)

try:
    h3 = doc_out.styles['Heading 3']
except KeyError:
    h3 = doc_out.styles.add_style('Heading 3', docx.enum.style.WD_STYLE_TYPE.PARAGRAPH)
h3.font.size = Pt(12)
h3.font.bold = True
h3.paragraph_format.line_spacing = 1.5
h3.paragraph_format.space_after = Pt(8)

normal = doc_out.styles['Normal']
normal.font.size = Pt(12)
normal.paragraph_format.line_spacing = 1.5
normal.paragraph_format.space_after = Pt(8)
normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

# Page Layout
for section in doc_out.sections:
    section.page_width = Inches(8.5)
    section.page_height = Inches(11.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

# ================= FRONT MATTER =================

# Cover Page
p = doc_out.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('[University Name]\nFaculty of [Faculty Name]\n\n\n')
run.font.size = Pt(16)
run.bold = True

run = p.add_run('"ShopEase Nepal: E-commerce Management System"\n\n\n')
run.font.size = Pt(18)
run.bold = True

run = p.add_run('A PROJECT REPORT\n\n\nSubmitted to\nDepartment of [Department Name]\n[College Name]\n\n')
run.font.size = Pt(14)
run.bold = True

run = p.add_run('In partial fulfillment of the requirements for the Bachelors in Computer Application\n\n\n')
run.font.italic = True

run = p.add_run('Submitted by\n\n')
run.bold = True

doc_out.add_paragraph('Anmol Chhetri (Lead)\nSahil Tuladhar\nSarang Limbu\nSanskriti Maharjan\nSmriti Tamang').alignment = WD_ALIGN_PARAGRAPH.CENTER

p2 = doc_out.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run('\n\nUnder the Supervision of\n[Supervisor Name]')
run2.bold = True

doc_out.add_page_break()

# Recommendation
doc_out.add_heading('Supervisor’s Recommendation', level=1)
doc_out.add_paragraph('I hereby recommend that this project prepared under my supervision by Anmol Chhetri, Sahil Tuladhar, Sarang Limbu, Sanskriti Maharjan, and Smriti Tamang entitled "ShopEase Nepal: E-commerce Management System" in the Partial Fulfillment of requirement for the degree of Bachelor in Computer Application is recommended for that final evaluation.')
doc_out.add_paragraph('\n\n_______________________\n[Supervisor Name]\nProject Supervisor\nBCA Department\n[College Name]')

doc_out.add_page_break()

# Approval
doc_out.add_heading('LETTER OF APPROVAL', level=1)
doc_out.add_paragraph('This is to certify that this project prepared by the team entitled "ShopEase Nepal" in partial fulfillment of the requirements for the degree of Bachelor in Computer Application has been evaluated. In our opinion it is satisfactory in the scope and quality as a project for the required degree.')
doc_out.add_paragraph('\n\n_______________________\n[Supervisor Name]\nSupervisor\n\n_______________________\n[Program Coordinator Name]\nProgram Coordinator\n\n_______________________\nInternal Examiner\n\n_______________________\nExternal Examiner')

doc_out.add_page_break()

# Abstract
doc_out.add_heading('ABSTRACT', level=1)
doc_out.add_paragraph('ShopEase Nepal is a comprehensive e-commerce platform tailored for the Nepalese market. All data is managed through a live Firebase Firestore database, with Cloudinary handling direct, secure image uploads. LocalStorage fallback mechanisms have been completely replaced by real-time cloud data. The application uses a live Firebase Firestore database and Google OAuth for authentication, cart, and product management.')

doc_out.add_page_break()

# Acknowledgement
doc_out.add_heading('ACKNOWLEDGEMENT', level=1)
doc_out.add_paragraph('Special thanks and gratitude are extended to the supervisor, [Supervisor Name], for providing the golden opportunity to undertake this project on the topic of ShopEase Nepal. This opportunity facilitated extensive research and the exploration of numerous new tools and technologies.')

doc_out.add_page_break()

doc_out.add_heading('Table of Contents', level=1)
doc_out.add_paragraph('Please Right-Click -> Update Field here to generate the TOC.')

doc_out.add_page_break()

doc_out.add_heading('List of Figures', level=1)
doc_out.add_paragraph('Figure 1: UI Screenshots... (See Appendix)')

# Chapters
doc_out.add_heading('Chapter 1\nINTRODUCTION', level=1)
doc_out.add_heading('1.1 Introduction', level=2)
doc_out.add_paragraph('ShopEase Nepal is a comprehensive e-commerce platform built to serve the Nepalese market. It features modern React context state management, a Firebase backend, and Cloudinary media integration to create a seamless shopping experience for users, while providing a powerful administrative dashboard for store owners.')

doc_out.add_heading('1.2 Problem Statement', level=2)
doc_out.add_paragraph('Traditional e-commerce platforms often lack tailored features for local markets like Nepal, such as precise regional delivery coverage mapping and localized payment integrations. Furthermore, many modern web applications struggle with efficient client-side state management without relying on heavy external libraries like Redux.')

doc_out.add_heading('1.3 Objectives', level=2)
doc_out.add_paragraph('1. To develop a responsive, modern e-commerce frontend using React 19 and Tailwind CSS.', style='List Bullet')
doc_out.add_paragraph('2. To implement secure multi-role access control (Admin, User) using Firebase Authentication.', style='List Bullet')
doc_out.add_paragraph('3. To provide an interactive map for checking delivery coverage across Nepalese provinces.', style='List Bullet')

doc_out.add_heading('1.4 Scope and Limitations', level=2)
doc_out.add_paragraph('Scope: The system covers product catalog browsing, cart management, user profiles, admin dashboards, and delivery mapping.', style='List Bullet')
doc_out.add_paragraph('Limitations: Real-time payment gateway integration (e-Sewa/Khalti) is currently simulated. Real-time chat support is not included.', style='List Bullet')

doc_out.add_heading('Chapter 2\nBACKGROUND STUDY AND LITERATURE REVIEW', level=1)
doc_out.add_heading('2.1 Study of existing system', level=2)
doc_out.add_paragraph('Existing systems like Daraz or Sastodeal offer comprehensive marketplaces, but often suffer from cluttered user interfaces and slow load times on mobile devices. ShopEase Nepal aims to provide a cleaner, more focused UI specifically optimized for performance using Vite and React.')

doc_out.add_heading('2.2 Literature Review', level=2)
doc_out.add_paragraph('Recent studies in web development highlight the shift towards lightweight state management (React Context) and utility-first CSS frameworks (Tailwind) to improve rendering speeds and developer experience. The use of NoSQL databases like Firebase Firestore has also been proven to accelerate the development of real-time applications.')

doc_out.add_heading('Chapter 3\nSYSTEM ANALYSIS AND DESIGN', level=1)

# Now extract from original document
doc_in = docx.Document('ShopEase_Documentation.docx')
copying = False
for p in doc_in.paragraphs:
    text = p.text.strip()
    if '4. System Architecture' in text or '2. System Requirements' in text or '5. State Management' in text or '8. Data Layer' in text or '9. User Roles' in text:
        doc_out.add_heading(text, level=2)
        copying = True
        continue
    elif text and text[0].isdigit() and text[1:3] == '. ':
        copying = False
        
    if copying and text:
        doc_out.add_paragraph(text, style=p.style.name)

doc_out.add_heading('Chapter 4\nIMPLEMENTATION AND TESTING', level=1)
copying = False
for p in doc_in.paragraphs:
    text = p.text.strip()
    if '3. Technology Stack' in text or '6. Pages and Features' in text or '7. Components' in text:
        doc_out.add_heading(text, level=2)
        copying = True
        continue
    elif text and text[0].isdigit() and text[1:3] == '. ':
        copying = False
        
    if copying and text:
        doc_out.add_paragraph(text, style=p.style.name)

doc_out.add_heading('Chapter 5\nFUTURE RECOMMENDATION', level=1)
doc_out.add_heading('5.1 Lesson Learnt / Outcome', level=2)
doc_out.add_paragraph('The development of ShopEase Nepal provided deep insights into modern React architecture, specifically the powerful combination of Context API and Tailwind CSS for rapid prototyping.')

doc_out.add_heading('5.2 Conclusion', level=2)
doc_out.add_paragraph('The system successfully meets its primary objectives, providing a robust, scalable e-commerce foundation that is ready for production deployment.')

doc_out.add_heading('5.3 Future Recommendations', level=2)
doc_out.add_paragraph('Future updates should include live payment gateway integration, AI-driven product recommendations, and mobile application versions using React Native.')

doc_out.save('ShopEase_Final_Report.docx')
print('Report generated successfully as ShopEase_Final_Report.docx')
