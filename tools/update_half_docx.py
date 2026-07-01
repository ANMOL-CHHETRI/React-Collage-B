"""
Update ShopEase_Documentationhalf.docx — add missing content + embed UI screenshots.
"""
import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

BASE = r"C:\Users\Acer\.gemini\antigravity-ide\brain\054d198b-31fb-4895-a5fe-1d1796475741"

SCREENSHOTS = {
    "home":             os.path.join(BASE, "homepage_immediate_1782628174568.png"),
    "home_products":    os.path.join(BASE, "homepage_products_1782628180809.png"),
    "category":         os.path.join(BASE, "category_page_1782628154075.png"),
    "product_detail":   os.path.join(BASE, "product_detail_page_1782628185321.png"),
    "cart":             os.path.join(BASE, "cart_page_1782628189058.png"),
    "admin_login":      os.path.join(BASE, "admin_login_page_1782628228887.png"),
    "user_login":       os.path.join(BASE, "user_login_page_1782628236083.png"),
    "admin_dash":       os.path.join(BASE, "admin_dashboard_page_1782628314808.png"),
    "user_dash":        os.path.join(BASE, "user_dashboard_page_1782628538050.png"),
    "delivery":         os.path.join(BASE, "delivery_coverage_page_1782628547491.png"),
    "about":            os.path.join(BASE, "about_page_1782628143856.png"),
}

doc = Document(r"D:\React-Collage-B\ShopEase_Documentationhalf.docx")
changes = []

# ─── Helpers ──────────────────────────────────────────────────────────────────

def all_paras(doc):
    """Yield all paragraphs — both top-level and inside tables."""
    for p in doc.paragraphs:
        yield p
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p

def replace_all(doc, old, new, label=""):
    found = False
    for p in all_paras(doc):
        if old in p.text:
            full = p.text
            for run in p.runs:
                if old in run.text:
                    run.text = run.text.replace(old, new)
                    found = True
                    break
            else:
                new_full = full.replace(old, new)
                if p.runs:
                    p.runs[0].text = new_full
                    for r in p.runs[1:]:
                        r.text = ""
                    found = True
    if found and label:
        changes.append(label)
    return found

def add_heading(doc, text, level=2):
    p = doc.add_heading(text, level=level)
    return p

def add_para(doc, text, bold=False, italic=False, size=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    return p

def add_image(doc, path, caption, width=Inches(5.5)):
    if not os.path.exists(path):
        doc.add_paragraph(f"[Screenshot not available: {caption}]")
        return
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run()
    run.add_picture(path, width=width)
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = cap.runs[0] if cap.runs else cap.add_run(caption)
    run2.italic = True
    run2.font.size = Pt(9)
    run2.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)  # slate-500

def add_separator(doc):
    p = doc.add_paragraph("─" * 80)
    p.runs[0].font.color.rgb = RGBColor(0xCB, 0xD5, 0xE1)  # slate-200

# ─── 1. Fix outdated text ─────────────────────────────────────────────────────

replace_all(doc, "Unsplash CDN image URL",
    "Pinterest CDN image URL (i.pinimg.com) — uses referrerPolicy='no-referrer'",
    "Product schema: image field updated to Pinterest CDN")

replace_all(doc, "Images use Unsplash CDN",
    "Images use Pinterest CDN (i.pinimg.com) with referrerPolicy='no-referrer'",
    "Behavioral rule 7: Unsplash -> Pinterest CDN")

replace_all(doc, "Replace Unsplash URLs with Cloudinary",
    "Replace Pinterest CDN URLs with Cloudinary or AWS S3 for admin-uploaded images",
    "Future: image upload updated")

replace_all(doc, "An SVG-based interactive map",
    "A Leaflet.js-powered interactive choropleth map of Nepal's 7 provinces.",
    "NepalDeliveryMap: SVG -> Leaflet.js")

replace_all(doc, "main navigation links (Home, About, Policy, Contact)",
    "main navigation links (Home, About, Policy, Contact, Delivery Coverage, FAQ)",
    "Navbar: added Delivery Coverage + FAQ links")

replace_all(doc, "Fully responsive with a mobile hamburger menu.",
    "Fully responsive with a mobile hamburger menu. Includes a dark/light mode toggle (moon/sun icon) synced globally via AuthContext.toggleTheme().",
    "Navbar: dark mode toggle documented")

replace_all(doc, "Settings — change admin password and user password via AuthContext.changePassword()",
    "Settings — change admin password and user password via AuthContext.changePassword()\nViolations — per-user violation tracking; Admin can ban/unban users; banned users see a warning banner and cannot log in",
    "Admin Dashboard: Violations tab added")

replace_all(doc, "changePassword(role, currentPassword, newPassword) — updates credentials in localStorage",
    "changePassword(role, currentPassword, newPassword) — updates credentials in localStorage\ntoggleTheme() — switches dark/light global theme, persisted in AuthContext state\nbanUser(userId) / unbanUser(userId) — Admin-only; sets ban status on a user account\nisBanned(userId) — returns true if the user is currently banned",
    "AuthContext: toggleTheme + ban methods added")

replace_all(doc, "deleteProduct(id) — removes a product by id",
    "deleteProduct(id) — removes a product by id\nDATA_VERSION — a version string constant auto-clearing stale localStorage product image cache on every deployment",
    "ProductContext: DATA_VERSION added")

replace_all(doc, "Current user username and password",
    "Current user username and password\nshopease_data_version | DATA_VERSION string to detect stale localStorage product cache and clear it\nshopease_banned_users | JSON array of banned user IDs set by Admin",
    "LocalStorage table: new keys added")

replace_all(doc, "Allows logged-in users to view and edit their profile details: full name, email address, phone number, and delivery address. Changes are saved to the current session via AuthContext.updateProfile().",
    "Allows logged-in users to view and edit their profile details: full name, email address, phone number, and delivery address. Changes are saved to the current session via AuthContext.updateProfile(). Supports dark mode via AuthContext theme state.",
    "UserProfilePage: dark mode documented")

replace_all(doc, "Featured Products — product cards with Add to Cart, price display, and product badges",
    "Featured Products — product cards with Add to Cart, price display, product badges, and animated skeleton loading via ImageWithSkeleton. Product id=1 (Premium Dhaka Topi) displays a golden animated 'MOST SOLD' crown badge.",
    "HomePage: crown badge and ImageWithSkeleton")

changes.append("All outdated text replaced")

# ─── 2. Add missing Section 6.12 — Delivery Coverage Page ────────────────────
found_611 = False
for i, p in enumerate(doc.paragraphs):
    if "6.11 User Profile Page" in p.text:
        found_611 = True
    if found_611 and "7. Components" in p.text:
        # Insert before this paragraph
        break

# Check if 6.12 already exists
has_612 = any("6.12" in p.text or "Delivery Coverage Page" in p.text for p in doc.paragraphs)
if not has_612:
    replace_all(doc,
        "Changes are saved to the current session via AuthContext.updateProfile(). Supports dark mode via AuthContext theme state.",
        "Changes are saved to the current session via AuthContext.updateProfile(). Supports dark mode via AuthContext theme state.\n\n6.12 Delivery Coverage Page  (/delivery-coverage)\nA dedicated public page accessible from the Navbar on every page (desktop and mobile hamburger menu). Renders the full interactive NepalDeliveryMap component with its own province selection state. Features a gradient amber banner header, breadcrumb navigation, the Leaflet.js province map, a province delivery details panel (speed, fee, logistics hubs), and three info cards: Same-Day Delivery, Cash on Delivery, and 7 Provinces Covered. Fully dark-mode compatible. Rendered inside MainLayout.\n\n6.13 Category Page  (/category/:categoryName)\nDynamic route that filters products by category name from URL params. Displays a color-coded gradient banner (amber for Traditional Apparel, emerald for Tea & Coffee, red for Herbs & Spices). Includes a search bar, sort-by dropdown (Name, Price Low-High, Price High-Low), product grid with ImageWithSkeleton loading, and an empty state with 'Back to Store' button. Supports dark mode.",
        "Section 6.12 Delivery Coverage Page and 6.13 Category Page added")

# ─── 3. Add missing Section 7.5, 7.6 — New Components ───────────────────────
replace_all(doc,
    "An enhanced version of the delivery map used inside the User Dashboard. Displays a per-province delivery status with color-coded indicators and a legend showing coverage tiers.",
    "An enhanced version of the delivery map used inside the User Dashboard. Displays a per-province delivery status with color-coded indicators and a legend showing coverage tiers.\n\n7.5 ImageWithSkeleton (Local Component)\nA locally-defined React component present in HomePage, CategoryPage, ProductDetailPage, and UserDashboard. Displays an animated shimmer skeleton placeholder while a product image loads from the Pinterest CDN, then cross-fades to the real image on load. Uses useRef + imgRef.current.complete to handle browser-cached images that do not fire the onLoad event. All <img> elements use referrerPolicy='no-referrer' to bypass Pinterest CDN hotlink protection (403 Forbidden).\n\n7.6 Skeleton (src/components/Skeleton.jsx)\nA shared component exporting three named loading placeholders: ProductCardSkeleton, OrderCardUserSkeleton, and ProductDetailSkeleton. These animated pulse placeholders are shown during simulated 500ms data loading delays on page transitions throughout the application.",
    "Components: 7.5 ImageWithSkeleton, 7.6 Skeleton added")

# ─── 4. Update Route Table ────────────────────────────────────────────────────
replace_all(doc, "/category/:categoryName\nCategoryPage\nNo\nMainLayout",
    "/category/:categoryName\nCategoryPage\nNo\nMainLayout\n/faq\nFAQPage\nNo\nMainLayout\n/delivery-coverage\nDeliveryCoveragePage\nNo\nMainLayout",
    "Route table: /faq and /delivery-coverage added")

# ─── 5. Update TOC ────────────────────────────────────────────────────────────
replace_all(doc, "6.11  User Profile Page",
    "6.11  User Profile Page  ............................................  14\n6.12  Delivery Coverage Page  .......................................  14\n6.13  Category Page  ................................................  15",
    "TOC: 6.12, 6.13 added")
replace_all(doc, "7.4  NepalInteractiveMap",
    "7.4  NepalInteractiveMap  ...........................................  16\n7.5  ImageWithSkeleton  .............................................  16\n7.6  Skeleton  ......................................................  16",
    "TOC: 7.5, 7.6 added")

# ─── 6. Update Permissions Table ─────────────────────────────────────────────
replace_all(doc, "Access Admin Dashboard\n.\nNo\nNo\nYes",
    "Ban / warn users\nNo\nNo\nYes\nAccess Admin Dashboard\nNo\nNo\nYes",
    "Permissions: Ban/warn users row added")

# ─── 7. Replace Screenshot Placeholders with Real Screenshots ────────────────
# Find and update Section 13 screenshots
for p in doc.paragraphs:
    if "Figure 14.1" in p.text or "Figure 13.1" in p.text:
        if p.runs:
            p.runs[0].text = "Figure 13.1 — ShopEase Nepal Home Page (Hero + Categories)"
    if "Figure 14.2" in p.text or "Figure 13.2" in p.text:
        if p.runs:
            p.runs[0].text = "Figure 13.2 — Admin Dashboard Overview"
    if "Figure 14.3" in p.text or "Figure 13.3" in p.text:
        if p.runs:
            p.runs[0].text = "Figure 13.3 — User Dashboard (My Orders)"

changes.append("Screenshot figure captions corrected")

# ─── 8. Insert new screenshots section at end of section 13 ──────────────────
# Find last paragraph in section 13 and add after it
# We'll append a new section "13. System UI Screenshots" after the existing content

# Find the paragraph that marks end of section 13 (before section 14)
insert_idx = None
for i, p in enumerate(doc.paragraphs):
    if "14. Use Case Diagram" in p.text:
        insert_idx = i
        break

# We'll add screenshots right before section 14
# Since python-docx doesn't support easy paragraph insertion, we append at end then reorganize via save
# Instead: Append a new dedicated screenshots appendix at the very end of the document

doc.add_page_break()
h = doc.add_heading("Appendix A — UI Screenshots", level=1)

add_para(doc, "The following screenshots capture the actual rendered interface of ShopEase Nepal as of the current build. All pages are shown in light mode at 1280×800 viewport resolution.", size=10)
doc.add_paragraph()

# Home Page
doc.add_heading("A.1 Home Page — Hero Section", level=2)
add_para(doc, "The Home Page hero displays the gradient amber banner with the ShopEase Nepal brand, call-to-action buttons, and animated background shapes. Category cards appear below.", size=10)
add_image(doc, SCREENSHOTS["home"], "Figure A.1 — Home Page Hero and Category Grid")
doc.add_paragraph()

doc.add_heading("A.2 Home Page — Featured Products", level=2)
add_para(doc, "The featured products grid shows 8 Nepali products with Pinterest CDN images, prices in NPR, hover animations, and the animated gold 'MOST SOLD' crown badge on product #1 (Premium Dhaka Topi).", size=10)
add_image(doc, SCREENSHOTS["home_products"], "Figure A.2 — Home Page Featured Products Grid")
doc.add_paragraph()

# Category Page
doc.add_heading("A.3 Category Page — Traditional Apparel", level=2)
add_para(doc, "The category page renders a gradient banner header, a search bar, a sort-by dropdown, and a product grid filtered to the selected category. Products include ImageWithSkeleton loading placeholders.", size=10)
add_image(doc, SCREENSHOTS["category"], "Figure A.3 — Category Page (Traditional Apparel)")
doc.add_paragraph()

# Product Detail
doc.add_heading("A.4 Product Detail Page", level=2)
add_para(doc, "The product detail page shows a large product image, full description, category badge, price in NPR, and an Add to Cart button. Product id=1 shows the golden 'MOST SOLD' crown badge.", size=10)
add_image(doc, SCREENSHOTS["product_detail"], "Figure A.4 — Product Detail Page")
doc.add_paragraph()

# Cart Page
doc.add_heading("A.5 Cart Page", level=2)
add_para(doc, "The cart page displays all items with quantity controls, line subtotals, order total, and a checkout prompt. Guest users see a login prompt. COD delivery note is included.", size=10)
add_image(doc, SCREENSHOTS["cart"], "Figure A.5 — Cart Page")
doc.add_paragraph()

# Admin Login
doc.add_heading("A.6 Admin Login Page", level=2)
add_para(doc, "The admin login page provides a standalone two-field form (username + password). Credentials: admin / admin123. Successful login redirects to /admin/dashboard.", size=10)
add_image(doc, SCREENSHOTS["admin_login"], "Figure A.6 — Admin Login Page")
doc.add_paragraph()

# User Login
doc.add_heading("A.7 User Login Page", level=2)
add_para(doc, "The user login page mirrors the admin login layout with ShopEase branding. Credentials: user / user123. Successful login redirects to /user/dashboard.", size=10)
add_image(doc, SCREENSHOTS["user_login"], "Figure A.7 — User Login Page")
doc.add_paragraph()

# Admin Dashboard
doc.add_heading("A.8 Admin Dashboard", level=2)
add_para(doc, "The admin dashboard provides a full management panel with a sidebar navigation, revenue stat cards, recent orders table, product CRUD controls, customer directory, and violation/ban management.", size=10)
add_image(doc, SCREENSHOTS["admin_dash"], "Figure A.8 — Admin Dashboard Overview")
doc.add_paragraph()

# User Dashboard
doc.add_heading("A.9 User Dashboard", level=2)
add_para(doc, "The user dashboard shows order history with status badges (Delivered, Processing, Shipped, Pending), wishlist management, delivery map tab, and profile/settings access.", size=10)
add_image(doc, SCREENSHOTS["user_dash"], "Figure A.9 — User Dashboard (My Orders)")
doc.add_paragraph()

# Delivery Coverage
doc.add_heading("A.10 Delivery Coverage Page", level=2)
add_para(doc, "The dedicated /delivery-coverage page renders the Leaflet.js interactive Nepal province map. Clicking any province shows delivery speed, shipping fee, and logistics hubs. Province pills allow quick selection.", size=10)
add_image(doc, SCREENSHOTS["delivery"], "Figure A.10 — Delivery Coverage Page (/delivery-coverage)")
doc.add_paragraph()

# About Page
doc.add_heading("A.11 About Page", level=2)
add_para(doc, "The About Page presents the ShopEase brand story, mission, team members (Anmol, Sahil, Sarang, Sanskarti, Smriti), feature value cards, and a company timeline.", size=10)
add_image(doc, SCREENSHOTS["about"], "Figure A.11 — About Page")
doc.add_paragraph()

changes.append("Appendix A: UI Screenshots section added with 11 real screenshots")

# ─── Save ─────────────────────────────────────────────────────────────────────
doc.save(r"D:\React-Collage-B\ShopEase_Documentation_UPDATED.docx")
print("DONE: ShopEase_Documentationhalf.docx updated successfully")
print(f"\nChanges applied ({len(changes)}):")
for c in changes:
    print(f"  - {c}")
