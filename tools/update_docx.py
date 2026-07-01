"""
Update ShopEase_Documentation.docx — full version that handles table cells too.
"""
import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Pt, RGBColor

doc = Document("ShopEase_Documentation.docx")

changes_made = []

def replace_in_para(para, old, new):
    """Replace text in a paragraph across all runs."""
    full = para.text
    if old not in full:
        return False
    for run in para.runs:
        if old in run.text:
            run.text = run.text.replace(old, new)
            return True
    # If split across runs, rebuild in first run
    if para.runs:
        para.runs[0].text = full.replace(old, new)
        for run in para.runs[1:]:
            run.text = ""
    return True

def search_all(doc, text):
    """Search both paragraphs and table cells, return list of (type, obj)."""
    results = []
    for para in doc.paragraphs:
        if text in para.text:
            results.append(('para', para))
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if text in para.text:
                        results.append(('cell', para))
    return results

def replace_all(doc, old, new, label=""):
    """Replace text in paragraphs and table cells throughout the document."""
    found = False
    for para in doc.paragraphs:
        if old in para.text:
            if replace_in_para(para, old, new):
                found = True
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if old in para.text:
                        if replace_in_para(para, old, new):
                            found = True
    if found and label:
        changes_made.append(label)
    return found

# ── 1. Image URL: Unsplash → Pinterest CDN ────────────────────────────────────
replace_all(doc,
    "Unsplash CDN image URL",
    "Pinterest CDN image URL (i.pinimg.com) — all images use referrerPolicy='no-referrer'",
    "Product schema: image field updated to Pinterest CDN"
)
replace_all(doc,
    "Replace Unsplash URLs with Cloudinary or AWS S3 uploads",
    "Replace Pinterest CDN URLs with Cloudinary or AWS S3 for admin-uploaded product images",
    "Future enhancement: image upload updated"
)
replace_all(doc,
    "Image Upload (Cloud)",
    "Image Upload (Cloud / Pinterest CDN Migration Complete)",
    "Future enhancement header updated"
)

# ── 2. Routing table: add /faq, /delivery-coverage, /category/:categoryName ───
replace_all(doc,
    "/category/:categoryName",
    "/category/:categoryName\n/faq\n/delivery-coverage",
    "Route table: added /faq and /delivery-coverage"
)

# ── 3. Navbar description ─────────────────────────────────────────────────────
replace_all(doc,
    "main navigation links (Home, About, Policy, Contact)",
    "main navigation links (Home, About, Policy, Contact, Delivery Coverage)",
    "Navbar: Delivery Coverage link documented"
)
replace_all(doc,
    "Fully responsive with a mobile hamburger menu.",
    "Fully responsive with a mobile hamburger menu. Includes a dark/light mode toggle (moon/sun icon) synced globally via AuthContext.toggleTheme().",
    "Navbar: dark mode toggle documented"
)

# ── 4. NepalDeliveryMap component ─────────────────────────────────────────────
replace_all(doc,
    "An SVG-based interactive map of Nepal rendered on the Home Page, showing delivery coverage zones across Nepal's seven provinces. Provinces highlight on hover and display coverage tooltips.",
    "A Leaflet.js-powered interactive choropleth map of Nepal's 7 provinces. Used on the Home Page and the dedicated /delivery-coverage page. Clicking a province displays delivery speed, shipping fee, and logistics hubs in real time. Built with GeoJSON data, animated orange marker pins, province quick-select pills, and full dark mode support.",
    "NepalDeliveryMap: description updated to Leaflet.js"
)

# ── 5. AuthContext methods ────────────────────────────────────────────────────
replace_all(doc,
    "changePassword(role, currentPassword, newPassword) — updates credentials in localStorage",
    "changePassword(role, currentPassword, newPassword) — updates credentials in localStorage\ntoggleTheme() — toggles dark/light global theme, persisted in AuthContext state\nbanUser(userId) / unbanUser(userId) — Admin-only; sets ban status on user accounts\nisBanned(userId) — checks if a user is currently banned",
    "AuthContext: new methods documented (theme, ban)"
)

# ── 6. Admin Dashboard features ───────────────────────────────────────────────
replace_all(doc,
    "Settings — change admin password and user password via AuthContext.changePassword()",
    "Settings — change admin password and user password via AuthContext.changePassword()\nViolations — view per-user violation count; ban/unban users; banned users see a warning banner and are blocked from logging in",
    "Admin Dashboard: violation/ban system documented"
)

# ── 7. ProductContext ─────────────────────────────────────────────────────────
replace_all(doc,
    "deleteProduct(id) — removes a product by id",
    "deleteProduct(id) — removes a product by id\nDATA_VERSION — a version string constant that auto-clears stale localStorage product cache on every version bump, ensuring users always load the latest product images after deployments",
    "ProductContext: DATA_VERSION documented"
)

# ── 8. LocalStorage table ─────────────────────────────────────────────────────
replace_all(doc,
    "shopease_user_credentials",
    "shopease_user_credentials",  # keep unchanged, we will add new row after
    ""
)
replace_all(doc,
    "Current user username and password",
    "Current user username and password\n\nshopease_data_version | DATA_VERSION string used to auto-clear stale product image cache on deploy\nshopease_banned_users | JSON array of banned user IDs set by Admin",
    "LocalStorage table: new keys added"
)

# ── 9. Permissions table ──────────────────────────────────────────────────────
replace_all(doc,
    "Access Admin Dashboard",
    "Ban / warn users",
    ""
)
# Revert bad replace
replace_all(doc,
    "Ban / warn users",
    "Ban / warn users\nAccess Admin Dashboard",
    "Permissions: Ban/warn users added before Access Admin Dashboard"
)

# ── 10. Pages section: User Profile description ───────────────────────────────
replace_all(doc,
    "Changes are saved to the current session via AuthContext.updateProfile().",
    "Changes are saved to the current session via AuthContext.updateProfile(). Supports dark mode via AuthContext theme state.",
    "UserProfilePage: dark mode documented"
)

# ── 11. Product Detail Page ───────────────────────────────────────────────────
replace_all(doc,
    "Includes breadcrumb navigation.",
    "Includes breadcrumb navigation. Products with id=1 display a special animated gold 'MOST SOLD' crown badge. Supports dark mode via AuthContext theme state. Product images use ImageWithSkeleton for graceful loading.",
    "ProductDetailPage: crown badge, dark mode, ImageWithSkeleton documented"
)

# ── 12. Home Page features list ───────────────────────────────────────────────
replace_all(doc,
    "Featured Products — product cards with Add to Cart, price display, and product badges",
    "Featured Products — product cards with Add to Cart, price display, product badges, and animated skeleton loading via ImageWithSkeleton. Product id=1 displays a golden 'MOST SOLD' crown badge.",
    "HomePage: crown badge and ImageWithSkeleton documented"
)
replace_all(doc,
    "Nepal Interactive Map — embedded delivery coverage visualization",
    "Nepal Interactive Map — embedded Leaflet.js delivery coverage map\nDark Mode — global dark/light theme toggle synced via AuthContext",
    "HomePage: dark mode feature documented"
)

# ── 13. Category Page — add new entry if not present ─────────────────────────
replace_all(doc,
    "6.12 Delivery Coverage Page",
    "6.12 Delivery Coverage Page",
    ""
)
# Add 6.12 after 6.11 description if not yet there
replace_all(doc,
    "Changes are saved to the current session via AuthContext.updateProfile(). Supports dark mode via AuthContext theme state.",
    "Changes are saved to the current session via AuthContext.updateProfile(). Supports dark mode via AuthContext theme state.\n\n6.12 Delivery Coverage Page (/delivery-coverage)\nA dedicated public page rendering the full NepalDeliveryMap component with province selection state. Accessible from the Navbar on every page (desktop + mobile). Includes a gradient amber header banner, breadcrumb navigation, the interactive Leaflet.js province map, a province delivery details panel, and three info cards: Same-Day Delivery, Cash on Delivery, and 7 Provinces Covered. Supports dark mode. Rendered inside MainLayout.",
    "Section 6.12 Delivery Coverage Page added"
)

# ── 14. Table of Contents — add new sections ──────────────────────────────────
replace_all(doc,
    "6.11  User Profile Page",
    "6.11  User Profile Page  ............................................  14\n6.12  Delivery Coverage Page  .......................................  14",
    "TOC: 6.12 Delivery Coverage added"
)
replace_all(doc,
    "7.4  NepalInteractiveMap",
    "7.4  NepalInteractiveMap  ...........................................  15\n7.5  ImageWithSkeleton  .............................................  16\n7.6  Skeleton  ......................................................  16",
    "TOC: 7.5 ImageWithSkeleton and 7.6 Skeleton added"
)

# ── 15. Components: add 7.5, 7.6 after 7.4 description ───────────────────────
replace_all(doc,
    "An enhanced version of the delivery map used inside the User Dashboard. Displays a per-province delivery status with color-coded indicators and a legend showing coverage tiers.",
    "An enhanced version of the delivery map used inside the User Dashboard. Displays a per-province delivery status with color-coded indicators and a legend showing coverage tiers.\n\n7.5 ImageWithSkeleton (Local Component)\nA locally-defined React component used inside HomePage, CategoryPage, ProductDetailPage, and UserDashboard. Displays an animated shimmer skeleton placeholder while a product image loads from the Pinterest CDN, then fades in the image on load. Uses useRef + imgRef.current.complete to handle browser-cached images that do not fire the onLoad event. All <img> tags include referrerPolicy='no-referrer' to bypass Pinterest CDN hotlink protection (403 Forbidden blocks).\n\n7.6 Skeleton (src/components/Skeleton.jsx)\nA shared component file exporting three named skeleton components: ProductCardSkeleton, OrderCardUserSkeleton, and ProductDetailSkeleton. These animated pulse-loading placeholders are shown during simulated data loading delays (setTimeout) on page transitions across the app.",
    "Components: 7.5 ImageWithSkeleton and 7.6 Skeleton added"
)

# ── Save ──────────────────────────────────────────────────────────────────────
doc.save("ShopEase_Documentation.docx")
print(f"DONE: ShopEase_Documentation.docx updated successfully")
print(f"\nChanges applied ({len(changes_made)}):")
for c in changes_made:
    print(f"  - {c}")
