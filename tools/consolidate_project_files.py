import os
import sys
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_BREAK

sys.stdout.reconfigure(encoding='utf-8')

# ─── 1. DEFINE LOG REPORTS DATA (WEEKS 1 - 8) ─────────────────────────
weeks_data = [
    {
        'week': 1,
        'phase': 'Project Initialization & Core Routing',
        'progress': '15% Complete',
        'tasks': [
            'Project Setup: Initialized React project with Vite, setting up the foundational environment.',
            'Routing Basics: Configured initial routes for HomePage and AboutPage.',
            'Navigation: Created basic Header and Navbar components.',
            'Cleanup: Removed unnecessary boilerplate code (index.css, old App.jsx contents).'
        ],
        'challenges': [
            'Configuration: Setting up initial project structure and dependencies.',
            'Boilerplate: Removing unwanted parts without breaking the Vite setup.'
        ],
        'solutions': [
            'Clean Slate: Established a clean starting point by methodically clearing out default Vite assets.',
            'Component Structure: Created a structured src/components directory from day one.'
        ],
        'contribs': [
            ('Anmol Chhetri (Lead)', 'Project Oversight', 'Project setup and initialization.'),
            ('Sahil Tuladhar', 'Backend Developer', 'Configured initial routing.'),
            ('Sarang Limbu', 'Frontend Developer', 'Created initial Header component.'),
            ('Sanskriti Maharjan', 'UI Component Developer', 'Basic CSS setup and cleanup.'),
            ('Smriti Tamang', 'QA & Testing', 'Initial project build verification.')
        ]
    },
    {
        'week': 2,
        'phase': 'Dashboard Layouts & Navigation',
        'progress': '30% Complete',
        'tasks': [
            'Dashboards: Developed basic UI scaffolding for Admin and User Dashboards.',
            'Navigation: Updated links and structure in Header/Navbar for better flow.',
            'Styling: Implemented responsive CSS updates across the application.',
            'Consistency: Standardized basic UI elements.'
        ],
        'challenges': [
            'Layout Consistency: Ensuring consistent styling across different dashboard views.',
            'Responsiveness: Making the initial Navbar work well on different screen sizes.'
        ],
        'solutions': [
            'Shared Classes: Centralized CSS classes and established a common layout structure.',
            'Flexbox/Grid: Utilized modern CSS layout techniques for responsive design.'
        ],
        'contribs': [
            ('Anmol Chhetri (Lead)', 'Project Oversight', 'Coordinated dashboard requirements.'),
            ('Sahil Tuladhar', 'Backend Developer', 'Admin Dashboard scaffolding.'),
            ('Sarang Limbu', 'Frontend Developer', 'Navbar enhancements and responsiveness.'),
            ('Sanskriti Maharjan', 'UI Component Developer', 'User Dashboard initial UI design.'),
            ('Smriti Tamang', 'QA & Testing', 'Navigation link testing.')
        ]
    },
    {
        'week': 3,
        'phase': 'Authentication UI & Documentation',
        'progress': '45% Complete',
        'tasks': [
            'Login UI: Enhanced User Login page UI with a visually appealing banner image.',
            'Documentation: Drafted initial project documentation and README files.',
            'Team Structure: Finalized development team roles and documented them.',
            'Forms: Set up basic input structures for authentication.'
        ],
        'challenges': [
            'Design Aesthetics: Designing an attractive login page that fits the premium e-commerce theme.',
            'Documentation: Ensuring all team roles and initial project specs were accurately captured.'
        ],
        'solutions': [
            'Visual Assets: Added a high-quality banner image and refined the form layout for better user experience.',
            'Markdown: Created structured markdown files to track progress and team assignments.'
        ],
        'contribs': [
            ('Anmol Chhetri (Lead)', 'Project Oversight', 'Drafted initial project documentation.'),
            ('Sahil Tuladhar', 'Backend Developer', 'Prepared forms for future login state management.'),
            ('Sarang Limbu', 'Frontend Developer', 'Implemented Login UI layout.'),
            ('Sanskriti Maharjan', 'UI Component Developer', 'Integrated banner and form styling.'),
            ('Smriti Tamang', 'QA & Testing', 'Form layout and responsiveness testing.')
        ]
    },
    {
        'week': 4,
        'phase': 'Theming, Maps, & QA Framework',
        'progress': '60% Complete',
        'tasks': [
            'Dark Mode: Implemented dark mode variants across the application.',
            'Map Integration: Added and fixed bugs related to the delivery coverage map.',
            'QA Setup: Added formal QA bug report templates and processes.',
            'FAQ Page: Implemented the FAQ Page UI.'
        ],
        'challenges': [
            'State Management: Ensuring dark mode applies consistently across all components.',
            'Version Control: Resolving merge conflicts resulting from parallel development tracks.'
        ],
        'solutions': [
            'Context API: Used React Context (AuthContext) to manage global theme state.',
            'Communication: Improved team coordination to reduce complex merge conflicts.'
        ],
        'contribs': [
            ('Anmol Chhetri (Lead)', 'Project Oversight', 'Merge conflict resolution and team coordination.'),
            ('Sahil Tuladhar', 'Backend Developer', 'Map functionality integration and fixes.'),
            ('Sarang Limbu', 'Frontend Developer', 'Developed the FAQ Page UI.'),
            ('Sanskriti Maharjan', 'UI Component Developer', 'Implemented global dark mode styles.'),
            ('Smriti Tamang', 'QA & Testing', 'Created QA bug report templates.')
        ]
    },
    {
        'week': 5,
        'phase': 'Product Catalog & User Profile',
        'progress': '75% Complete',
        'tasks': [
            'Product Data: Added Traditional Apparel products and improved sorting logic.',
            'User Profile: Redesigned the User Profile UI for better usability.',
            'Footer: Created a global footer component and added it to multiple pages.',
            'Image Optimization: Addressed image loading and referrer policy issues.'
        ],
        'challenges': [
            'Security & Performance: Preventing image referrer leakage and optimizing image loading from external CDNs.',
            'Component Reusability: Building a footer that works consistently across all page layouts.'
        ],
        'solutions': [
            'Security Policies: Added referrerPolicy=\'no-referrer\' to img tags.',
            'UX Enhancements: Created an ImageWithSkeleton component for smoother image loading experiences.'
        ],
        'contribs': [
            ('Anmol Chhetri (Lead)', 'Project Oversight', 'Product data management and curation.'),
            ('Sahil Tuladhar', 'Backend Developer', 'Implemented robust product sorting logic.'),
            ('Sarang Limbu', 'Frontend Developer', 'Created and integrated the global footer component.'),
            ('Sanskriti Maharjan', 'UI Component Developer', 'Redesigned the User Profile interface.'),
            ('Smriti Tamang', 'QA & Testing', 'Verified image loading behaviors and skeleton animations.')
        ]
    },
    {
        'week': 6,
        'phase': 'Page Refinements & E2E Testing Prep',
        'progress': '85% Complete',
        'tasks': [
            'Page Updates: Refined Contact, UserLogin, ProductDetails, and AdminLogin pages.',
            'Shopping Experience: Enhanced overall user shopping experience and order action flows.',
            'Testing: Added initial Playwright configuration for End-to-End (E2E) testing.',
            'Cleanup: Removed unnecessary images and files from the repository.'
        ],
        'challenges': [
            'Complex Flows: Handling complex user flows and ensuring all pages meet high quality standards.',
            'Testing Setup: Configuring Playwright correctly for a React/Vite environment.'
        ],
        'solutions': [
            'Iterative Polish: Conducted thorough manual testing and iterative UI refinements.',
            'Test Infrastructure: Established the foundational setup for automated E2E tests.'
        ],
        'contribs': [
            ('Anmol Chhetri (Lead)', 'Project Oversight', 'Enhanced order actions and user shopping flows.'),
            ('Sahil Tuladhar', 'Backend Developer', 'Set up initial Playwright configuration.'),
            ('Sarang Limbu', 'Frontend Developer', 'Updated ProductDetails page functionality.'),
            ('Sanskriti Maharjan', 'UI Component Developer', 'Refined Contact page UI.'),
            ('Smriti Tamang', 'QA & Testing', 'E2E test planning and manual page verification.')
        ]
    },
    {
        'week': 7,
        'phase': 'Final UI/UX Polish & Authentication Flows',
        'progress': '95% Complete',
        'tasks': [
            'Sign-up UI Redesign: Implemented a new slide-up modal drawer for user registration.',
            'Login Animation: Added a cinematic store opening animation overlay on login.',
            'Admin Login Revamp: Refactored Admin Login to a clean card layout with no images.',
            'Hidden Access: Created a secret 4-corner tap sequence to unlock admin login securely.',
            'Auth Flow Management: Patched AuthContext to give UserLogin page control over redirect timing.'
        ],
        'challenges': [
            'Animation Sync: Ensuring complex CSS keyframes timed perfectly with React state and route redirects.',
            'Security via Obscurity: Implementing a hidden admin trigger without breaking normal user flows.'
        ],
        'solutions': [
            'Controlled Timeouts: Used setTimeout linked to exact CSS animation durations for routing.',
            'State Separation: Handled sign-up flow through local states inside UserLoginPage rather than global context.'
        ],
        'contribs': [
            ('Anmol Chhetri (Lead)', 'Project Oversight', 'Managed UI/UX transition and reviewed auth workflow logic.'),
            ('Sahil Tuladhar', 'Backend Developer', 'Handled localStorage data structures and state persistence.'),
            ('Sarang Limbu', 'Frontend Developer', 'Developed core React state logic for sign-up modal and Easter eggs.'),
            ('Sanskriti Maharjan', 'UI Component Developer', 'Designed and implemented custom CSS keyframe animations.'),
            ('Smriti Tamang', 'QA & Testing', 'Verified visual timing of animations and tested the secret admin trigger.')
        ]
    },
    {
        'week': 8,
        'phase': 'Coworking Space & UI Polish',
        'progress': '100% Complete',
        'tasks': [
            'Coworking Space: Designed and implemented a minimal, high-end Coworking Space landing page (/coworking) with space tour modals and specs.',
            'Cart Booking: Integrated Coworking Pass bookings (Daily, Monthly, and Team Cabin plans) directly into the shopping cart.',
            'Payment Badges: Replaced static payment placeholders (COD, eSewa, Khalti) in the footer with fully functional hover links.',
            'Text Polish: Fixed navigation layout alignments on mobile viewports by converting "Log in" -> "Login" and "Sign up" -> "Signup" to prevent text wrapping.'
        ],
        'challenges': [
            'Visual Balance: Achieving a high-end corporate architectural site layout without using clunky AI-generated templates.',
            'Seamless Sinks: Wiring up custom non-product items (coworking bookings) to flow seamlessly through the global shopping cart state.'
        ],
        'solutions': [
            'Minimalist Design: Replaced flashing background neon gradient shapes with precise borders, substantial whitespace, and solid neutral colors.',
            'Dynamic Schema Mapping: Structured coworking plans to conform to the primary products schema to leverage standard checkout and cart procedures.'
        ],
        'contribs': [
            ('Anmol Chhetri (Lead)', 'Project Oversight', 'Managed clean layout aesthetic guidelines and approved booking integrations.'),
            ('Sahil Tuladhar', 'Backend Developer', 'Persisted booking items and mapped local storage schemas.'),
            ('Sarang Limbu', 'Frontend Developer', 'Designed layout preview modals and mobile menu adjustments.'),
            ('Sanskriti Maharjan', 'UI Component Developer', 'Coded responsive consultation schedule request forms.'),
            ('Smriti Tamang', 'QA & Testing', 'Executed browser layout consistency checks on multiple screen configurations.')
        ]
    }
]

# ─── 2. GENERATE SINGLE WEEKLY_PROGRESS_LOG.MD ────────────────────────
print("Generating WEEKLY_PROGRESS_LOG.md...")
md_lines = ["# ShopEase Nepal — Weekly Progress Logs\n"]
for w in weeks_data:
    md_lines.append(f"## Week {w['week']}: {w['phase']}")
    md_lines.append(f"**Status:** {w['progress']}\n")
    
    md_lines.append("### Tasks Completed")
    for task in w['tasks']:
        md_lines.append(f"- {task}")
    md_lines.append("")
    
    md_lines.append("### Challenges")
    for c in w['challenges']:
        md_lines.append(f"- {c}")
    md_lines.append("")
    
    md_lines.append("### Solutions")
    for s in w['solutions']:
        md_lines.append(f"- {s}")
    md_lines.append("")
    
    md_lines.append("### Team Contributions")
    md_lines.append("| Member | Focus Area | Deliverables |")
    md_lines.append("| :--- | :--- | :--- |")
    for name, focus, deliv in w['contribs']:
        md_lines.append(f"| **{name}** | {focus} | {deliv} |")
    md_lines.append("\n---\n")

# Append remaining tasks
md_lines.append("## Project Outlook & Remaining Tasks\n")
md_lines.append("While the frontend and UI/UX flows are 100% complete, there are a few critical milestones left before the final launch of ShopEase Nepal:\n")
md_lines.append("### 1. Database Architecture & Setup")
md_lines.append("- Designing and provisioning the primary relational database (e.g., MySQL or PostgreSQL) to store users, products, orders, and wishlist data.")
md_lines.append("- Establishing secure connections and data models mapping directly to the frontend schemas.\n")
md_lines.append("### 2. Backend REST API Integration")
md_lines.append("- Currently, the application relies on mock static data and local storage for authentication, cart, and product management. The next major phase will connect the frontend to a live backend database.")
md_lines.append("- Endpoints to integrate include: User Auth, Admin Product CRUD, Order Processing, and Wishlist syncing.\n")
md_lines.append("### 3. Payment Gateway Setup")
md_lines.append("- Integrating Khalti and e-Sewa payment APIs to process real transactions during the checkout flow.\n")
md_lines.append("### 4. Production Deployment")
md_lines.append("- Configuring GitHub Pages or a dedicated hosting environment (like Vercel/Netlify) for public access.")
md_lines.append("- Setting up automated CI/CD pipelines via GitHub Actions for continuous deployment.\n")

with open("WEEKLY_PROGRESS_LOG.md", "w", encoding="utf-8") as f:
    f.write("\n".join(md_lines))
print("Successfully generated WEEKLY_PROGRESS_LOG.md")

# ─── 3. GENERATE SINGLE WEEKLY_PROGRESS_LOG.DOCX ──────────────────────
print("Generating WEEKLY_PROGRESS_LOG.docx...")
log_doc = Document()
for idx, w in enumerate(weeks_data):
    log_doc.add_heading(f"WEEKLY PROGRESS LOG - WEEK {w['week']}", 0)
    
    p = log_doc.add_paragraph()
    p.add_run('Project Title: ').bold = True
    p.add_run('ShopEase Nepal (Ecommerce Platform)\n')
    p.add_run('Current Phase: ').bold = True
    p.add_run(f"{w['phase']}\n")
    p.add_run('Overall Progress: ').bold = True
    p.add_run(w['progress'])

    log_doc.add_heading('1. TASKS COMPLETED', level=1)
    for task in w['tasks']:
        log_doc.add_paragraph(task, style='List Bullet')

    log_doc.add_heading('2. CHALLENGES', level=1)
    for c in w['challenges']:
        log_doc.add_paragraph(c, style='List Bullet')

    log_doc.add_heading('3. SOLUTIONS', level=1)
    for s in w['solutions']:
        log_doc.add_paragraph(s, style='List Bullet')

    log_doc.add_heading('4. INDIVIDUAL TASK CONTRIBUTIONS', level=1)
    table = log_doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Member'
    hdr_cells[1].text = 'Focus Area'
    hdr_cells[2].text = 'Key Deliverables'

    for item in w['contribs']:
        row_cells = table.add_row().cells
        row_cells[0].text = item[0]
        row_cells[1].text = item[1]
        row_cells[2].text = item[2]

    log_doc.add_heading('5. LEADERSHIP REVIEW & STATUS REPORT', level=1)
    if w['week'] < 8:
        log_doc.add_paragraph(f"The project successfully completed Week {w['week']}, reaching {w['progress']}. The team focused heavily on {w['phase'].lower()}, resolving key challenges along the way.")
    else:
        log_doc.add_paragraph('The project has successfully reached 100% completion. With core pages, checkout states, and the interactive Nepalese delivery map finalized, the platform now features a highly clean, premium workspace pass booking landing page integrated with the shopping cart, making it ready for a backend database hookup.')

    p2 = log_doc.add_paragraph()
    p2.add_run('Current Status: ').bold = True
    p2.add_run(f"{w['progress']}\n")
    p2.add_run('Next Milestone: ').bold = True
    if w['week'] < 8:
        p2.add_run(f"Proceed to Week {w['week']+1} objectives.")
    else:
        p2.add_run('Full database deployment and backend API gateway wiring.')

    if idx < len(weeks_data) - 1:
        log_doc.add_page_break()

# Add remaining tasks section to the end of log docx
log_doc.add_page_break()
log_doc.add_heading('PROJECT OUTLOOK & REMAINING TASKS', 0)
log_doc.add_paragraph('While the frontend and UI/UX flows are 100% complete, there are a few critical milestones left before the final launch of ShopEase Nepal:')

log_doc.add_heading('1. Database Architecture & Setup', level=1)
log_doc.add_paragraph('Designing and provisioning the primary relational database (e.g., MySQL or PostgreSQL) to store users, products, orders, and wishlist data.', style='List Bullet')
log_doc.add_paragraph('Establishing secure connections and data models mapping directly to the frontend schemas.', style='List Bullet')

log_doc.add_heading('2. Backend REST API Integration', level=1)
log_doc.add_paragraph('Currently, the application relies on mock static data and local storage for authentication, cart, and product management. The next major phase will connect the frontend to a live backend database.', style='List Bullet')
log_doc.add_paragraph('Endpoints to integrate include: User Auth, Admin Product CRUD, Order Processing, and Wishlist syncing.', style='List Bullet')

log_doc.add_heading('3. Payment Gateway Setup', level=1)
log_doc.add_paragraph('Integrating Khalti and e-Sewa payment APIs to process real transactions during the checkout flow.', style='List Bullet')

log_doc.add_heading('4. Production Deployment', level=1)
log_doc.add_paragraph('Configuring GitHub Pages or a dedicated hosting environment (like Vercel/Netlify) for public access.', style='List Bullet')
log_doc.add_paragraph('Setting up automated CI/CD pipelines via GitHub Actions for continuous deployment.', style='List Bullet')

log_doc.save("WEEKLY_PROGRESS_LOG.docx")
print("Successfully generated WEEKLY_PROGRESS_LOG.docx")

# ─── 4. UPDATE SHOPEASE_DOCUMENTATION.DOCX ───────────────────────────
print("Updating ShopEase_Documentation.docx...")

# Load the file with screenshots (updated) and copy it back as ShopEase_Documentation.docx
source_doc_path = "ShopEase_Documentation_UPDATED.docx"
if not os.path.exists(source_doc_path):
    source_doc_path = "ShopEase_Documentation.docx"

doc = Document(source_doc_path)

def replace_in_para(para, old, new):
    full = para.text
    if old not in full:
        return False
    for run in para.runs:
        if old in run.text:
            run.text = run.text.replace(old, new)
            return True
    if para.runs:
        para.runs[0].text = full.replace(old, new)
        for run in para.runs[1:]:
            run.text = ""
    return True

def replace_all(doc, old, new):
    found = False
    for para in doc.paragraphs:
        if old in para.text:
            if replace_in_para(para, old, new):
                found = True
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if old in para.text:
                        if replace_in_para(para, old, new):
                            found = True
    return found

# 4.1 Update Table of Contents
replace_all(doc, 
            "6.12  Delivery Coverage Page  .......................................  14", 
            "6.12  Delivery Coverage Page  .......................................  14\n6.13  Coworking Space Page  .........................................  14")

# 4.2 Update Routing Table in section 4.3 or similar (replace /delivery-coverage with /delivery-coverage and /coworking)
replace_all(doc, 
            "/delivery-coverage", 
            "/delivery-coverage\n/coworking")

# 4.3 Update Section 6.12 block to append 6.13 block
old_delivery_desc = "6.12 Delivery Coverage Page (/delivery-coverage)\nA dedicated public page rendering the full NepalDeliveryMap component with province selection state. Accessible from the Navbar on every page (desktop + mobile). Includes a gradient amber header banner, breadcrumb navigation, the interactive Leaflet.js province map, a province delivery details panel, and three info cards: Same-Day Delivery, Cash on Delivery, and 7 Provinces Covered. Supports dark mode. Rendered inside MainLayout."
new_delivery_desc_with_coworking = old_delivery_desc + "\n\n6.13 Coworking Space Page (/coworking)\nA premium, clean workspace landing page outlining hot desking, meeting boardrooms, and private team suites. Features modular layout layouts, interactive space previews with detail modals, a productive amenities index (WiFi, backup power, estate coffee), and a consultation tour scheduling form. Includes e-commerce shopping cart integration to book day passes and monthly subscriptions directly."

replace_all(doc, old_delivery_desc, new_delivery_desc_with_coworking)

# Save as the unified main documentation
doc.save("ShopEase_Documentation.docx")
print("Successfully saved updated ShopEase_Documentation.docx")

# ─── 5. CLEAN UP ROUGE LOGS & COPIES ─────────────────────────────────
print("Cleaning up redundant log and document duplicates...")
files_to_delete = [
    "WEEKLY_PROGRESS_LOG_WEEK_1.docx",
    "WEEKLY_PROGRESS_LOG_WEEK_2.docx",
    "WEEKLY_PROGRESS_LOG_WEEK_3.docx",
    "WEEKLY_PROGRESS_LOG_WEEK_4.docx",
    "WEEKLY_PROGRESS_LOG_WEEK_5.docx",
    "WEEKLY_PROGRESS_LOG_WEEK_6.docx",
    "WEEKLY_PROGRESS_LOG_WEEK_7.docx",
    "WEEKLY_PROGRESS_LOG_ALL_WEEKS.docx",
    "WEEKLY_PROGRESS_LOG_ALL_WEEKS_FINAL.docx",
    "WEEKLY_PROGRESS_LOG_ALL_WEEKS_FINAL_V2.docx",
    "ShopEase_Documentationhalf.docx",
    "ShopEase_Documentation_UPDATED.docx",
    "generate_combined_log.py",
    "generate_logs.py"
]

for f in files_to_delete:
    if os.path.exists(f):
        try:
            os.remove(f)
            print(f"Deleted redundant file: {f}")
        except Exception as e:
            print(f"Error deleting {f}: {e}")

print("\nAll project files consolidated successfully!")
